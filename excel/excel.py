import openpyxl
from openpyxl.styles import Color, PatternFill, Font, Border
from openpyxl.styles import colors
from openpyxl.cell import Cell
import win32com.client
import os
import docx
from docx import Document
import datetime
import  smtplib
from win32com.client import Dispatch
excel = win32com.client.Dispatch("Excel.Application")
wb = openpyxl.load_workbook('students.xlsx')

print(wb.sheetnames)

sheet_1 = wb.active

# for i in range(1, sheet_1.max_row):
# print(sheet_1.cell(row=i, column=2).value)

# ищу группы
groups_list = []
groups_id = []
prevGroup = []
data = sheet_1.cell(row=2, column=3).value
for i in range(2, sheet_1.max_row):
    if str(sheet_1.cell(row=i, column=3).value) not in groups_list:
        groups_list.append(str(sheet_1.cell(row=i, column=3).value))
        groups_id.append(i)

print(groups_list)
print(groups_id)

ws1 = wb.create_sheet("5.01")

ws2 = wb.create_sheet("5.02")
ws3 = wb.create_sheet("5.03")
for i in range(1, 7):
    ws1.cell(row=1, column=i).value = sheet_1.cell(row=1, column=i).value
for i in range(1, 7):
    ws2.cell(row=1, column=i).value = sheet_1.cell(row=1, column=i).value
for i in range(1, 7):
    ws3.cell(row=1, column=i).value = sheet_1.cell(row=1, column=i).value

wb.save("students_edit.xlsx")
for i in range(2, groups_id[1] - 1):
    ws1.cell(row=i, column=1).value = i - 1

for i in range(1, int(groups_id[1]) - 1):
    for j in range(2, 7):
        ws1.cell(row=i, column=j).value = sheet_1.cell(row=i, column=j).value
wb.save("students_edit.xlsx")

for i in range(2, groups_id[1] - 1):
    ws2.cell(row=i, column=1).value = i - 1

for i in range(int(groups_id[1]), int(groups_id[2])):
    for j in range(2, 7):
        ws2.cell(row=i - 6, column=j).value = sheet_1.cell(row=i, column=j).value
wb.save("students_edit.xlsx")

for i in range(2, groups_id[1] - 1):
    ws3.cell(row=i, column=1).value = i - 1

for i in range(int(groups_id[2]), sheet_1.max_row):
    for j in range(2, 7):
        ws3.cell(row=i - 11, column=j).value = sheet_1.cell(row=i, column=j).value

ws1.cell(row=1, column=7).value = "Mark"
ws2.cell(row=1, column=7).value = "Mark"
ws3.cell(row=1, column=7).value = "Mark"

value = 0
for i in range(2, ws1.max_row + 1):

    value = 0
    for j in range(4, 7):
        value = value + ws1.cell(row=i, column=j).value
        print(value)
        ws1.cell(row=i, column=7).value = round(value / 3)

for i in range(2, ws2.max_row + 1):
    value = 0
    for j in range(4, 7):
        value = value + ws2.cell(row=i, column=j).value
        ws2.cell(row=i, column=7).value = round(value / 3)

for i in range(2, ws3.max_row + 1):
    value = 0
    for j in range(4, 7):
        value = value + ws3.cell(row=i, column=j).value
        ws3.cell(row=i, column=7).value = round(value / 3)

wb.save("students_edit.xlsx")

print("SORTING...")
wb.close()
temp = 0
#for i in range(2, ws3.max_row + 1):
 #   temp = 0
  #  if int(ws3.cell(row=i, column=7).value) > int(ws3.cell(row=i + 1, column=7).value):
   #     temp = int(ws3.cell(row=i + 1, column=7).value)
    #    ws3.cell(row=i + 1, column=7).value = ws3.cell(row=i, column=7).value
    #    ws3.cell(row=i, column=7).value = temp
     #   ws3.cell(
#ws2.auto_filter.ref = ws2.dimensions
#ws3.auto_filter.ref = ws3.dimensions



#ws1.auto_filter.ref = "A2:G6"
#ws1.auto_filter.add_filter_column(0, ["Kiwi", "Apple", "Mango"])
#ws1.auto_filter.add_sort_condition("B2:B15")


wb = excel.Workbooks.Open(os.path.join(os.getcwd(), "students_edit.xlsx"))
ws1 = wb.Worksheets("5.01")
ws1.Range('A2:G6').Sort(Key1=ws1.Range('G6'), Order1=2, Orientation=1)
ws2 = wb.Worksheets("5.02")
ws2.Range('A2:G6').Sort(Key1=ws2.Range('G6'), Order1=2, Orientation=1)
ws3 = wb.Worksheets("5.03")
ws3.Range('A2:G6').Sort(Key1=ws3.Range('G6'), Order1=2, Orientation=1)

wb.Save()
excel.Application.Quit()

redFill = PatternFill(start_color='FFFF0000', end_color='FFFF0000', fill_type='solid')

wb = openpyxl.load_workbook('students_edit.xlsx')
ws1 = wb["5.01"]
ws2 = wb["5.02"]
ws3 = wb["5.03"]
dumbs = []
for i in range(2, ws1.max_row + 1):
    if ws1.cell(row=i, column=7).value <= 60:
        ws1.cell(row=i, column=7).fill = redFill
        dumbs.append(ws1.cell(row=i, column=2).value)
for i in range(2, ws2.max_row + 1):
    if ws2.cell(row=i, column=7).value <= 60:
        ws2.cell(row=i, column=7).fill = redFill
        dumbs.append(ws2.cell(row=i, column=2).value)
for i in range(2, ws3.max_row + 1):
    if ws3.cell(row=i, column=7).value <= 60:
        ws3.cell(row=i, column=7).fill = redFill
        dumbs.append(ws3.cell(row=i, column=2).value)

wb.save("students_edit.xlsx")

ws4 = wb.create_sheet("Borjniki")
cntr = 0
ws4.cell(row=1, column=1).value = "Borjniki"
for i in range(2, len(dumbs) + 2):
    ws4.cell(row=i, column=1).value = dumbs[cntr]
    cntr = cntr+1

wb.save("students_edit.xlsx")
print(dumbs)

document = Document()
document.add_heading("Люди - позор академії, на відрахування")

cntr = 0
for i in range(0,len(dumbs)):
    paragraph = document.add_paragraph(dumbs[i])

document.add_heading(str(datetime.datetime.now()))

document.save("borjniki.docx")
print("Message sendibg")
#pochta
to='////'
fromname='Dmitriy TE'
fromemail='///'
subject='Тест программы'
body='Файл отчислений'

message=''
message+= "To: " + to + "\n"
message+= "From: \"" + fromname + "\" <" + fromemail + ">\n"
message+= "Subject: " + subject + "\n"
message+= "\n"
message+= body

mailserver = smtplib.SMTP('smtp.rambler.ru', 465)
mailserver.ehlo()
mailserver.starttls()
mailserver.ehlo()  #again
mailserver.login('pythonte///11@rambler.ru', '///')
mailserver.sendmail(fromemail, to, message)
mailserver.quit()                                 # Выходим
print("message sent")